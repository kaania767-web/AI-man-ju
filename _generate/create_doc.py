"""生成部署记录 Word 文档"""
import os
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# 样式
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# ── 封面 ──
title = doc.add_heading('《浪漫主义流浪家》', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0xE8, 0xA0, 0x40)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('AI 漫剧 · 从本地到网络的完整部署记录')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6B, 0x4C, 0x6E)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('作者：唐文浚　　部署日期：2026年6月24日')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════ 第一章 ══════════
doc.add_heading('第一章：项目原本在哪里？', level=1)

doc.add_paragraph(
    '《浪漫主义流浪家》是一个完全由 AI 生成的漫剧作品。'
    '从剧本创作、分镜脚本、AI 绘图，到最终合成视频——'
    '所有文件都存放在你电脑的这个文件夹里：'
)
p = doc.add_paragraph()
run = p.add_run(r'C:\Users\唐文浚\Documents\AI漫剧')
run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
run.bold = True

doc.add_paragraph(
    '里面有 30 多个文件，包括 28 张 AI 生成的漫画分镜图、'
    '一个 95MB 的高清视频、一整套 Python 生成脚本、'
    '以及一个 HTML 格式的观看页面。'
    '\n\n问题在于——这个页面只有你自己能打开。别人想看，'
    '就需要一个「网络链接」。'
)

# ══════════ 第二章 ══════════
doc.add_heading('第二章：为什么别人访问不了？', level=1)

doc.add_paragraph(
    '普通的 HTML 文件是跑在你自己电脑上的。就像一本纸质书放在你家书架上——'
    '客人要来看，必须走进你家才行。'
)
doc.add_paragraph(
    '要让全世界的人都能看，需要做两件事：'
)

items = [
    ('压缩视频', '原始视频 95MB，太大，传不到网上去。就像寄快递——太重的包裹寄不了。'),
    ('部署到服务器', '把文件上传到一个 24 小时开机的网络服务器上，'
     '它会替你把页面发给每一个访问的人。'),
]
for title_text, desc in items:
    p = doc.add_paragraph()
    run = p.add_run('► ' + title_text + '：')
    run.bold = True
    p.add_run(desc)

# ══════════ 第三章 ══════════
doc.add_heading('第三章：压缩视频——从 95MB 到 8MB', level=1)

doc.add_paragraph(
    '用了一个叫 ffmpeg 的工具来处理。它是视频处理领域的"瑞士军刀"，'
    '我们用它的三个功能来瘦身：'
)

steps = [
    ('缩小分辨率', '从 1920x1080 降到 1280x720。'
     '画面依然清晰，但数据量少了一半多。'),
    ('降低码率', '用 CRF 28 的参数重新编码。'
     '这是一种"智能压缩"模式——对画面简单的部分压得多，复杂的部分压得少。'),
    ('优化网络播放', '加了 faststart 参数，'
     '让视频在浏览器里可以"边下边播"，不用等全部下载完。'),
]
for i, (title_text, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {title_text}：')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    '结果：视频从 95MB 降到了 8.3MB（约原来的 1/11），画质仍然很好。'
)

# ══════════ 第四章 ══════════
doc.add_heading('第四章：打包站点', level=1)

doc.add_paragraph(
    '把所有需要上线的文件打包成一个 ZIP，就像把你的书架装箱托运：'
)

items_bullet = [
    ('index.html', '观看页面本身'),
    ('output/ 文件夹', '28 张漫画图片 + 压缩后的视频'),
    ('netlify.toml', '服务器配置（告诉服务器：视频要大缓存）'),
]
for name, desc in items_bullet:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name)
    run.bold = True
    p.add_run(' — ' + desc)

doc.add_paragraph('打包后总大小 27MB——网上随便传。')

# ══════════ 第五章 ══════════
doc.add_heading('第五章：部署到 Netlify——让全世界能访问', level=1)

doc.add_paragraph(
    'Netlify 是一个专门托管网站的服务。它就像一间"云上展厅"——'
    '你把作品交给它，它帮你 24 小时开门迎客。'
)

doc.add_paragraph('部署过程极其简单：')

steps2 = [
    ('打开网站', ' https://app.netlify.com/drop'),
    ('拖拽文件', '把 ZIP 包拖进浏览器窗口'),
    ('放手，等 7 秒', 'Netlify 自动解压、配置、上线'),
    ('拿到链接', ' https://gregarious-paletas-1af7bf.netlify.app'),
]
for i, (title_text, desc) in enumerate(steps2, 1):
    p = doc.add_paragraph()
    run = p.add_run(str(i) + '. ' + title_text)
    run.bold = True
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run(desc)

doc.add_paragraph()

# ══════════ 第六章 ══════════
doc.add_heading('第六章：最终成果', level=1)

p = doc.add_paragraph()
run = p.add_run('你的漫剧现在有了一个公开链接：')
run.bold = True

p2 = doc.add_paragraph()
run2 = p2.add_run('https://gregarious-paletas-1af7bf.netlify.app')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
run2.bold = True

doc.add_paragraph()

# 数据表格
table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.columns[0].width = Cm(5)
table.columns[1].width = Cm(10)

data = [
    ('指标', '数值'),
    ('原始视频大小', '95 MB（1920x1080）'),
    ('优化后视频大小', '8.3 MB（1280x720）'),
    ('站点总大小', '27 MB'),
    ('部署用时', '7 秒'),
    ('链接可用性', '永久在线，全球可访问'),
]
for i, (k, v) in enumerate(data):
    row = table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_paragraph(
    '从此，任何人只要点开这个链接，就能看到你的 AI 漫剧作品——'
    '在手机上、平板上、电脑上，无论他们在哪个城市。'
)

# 结尾
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 完 ——')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)

# 保存
output_path = r'C:\Users\唐文浚\Documents\AI漫剧\部署记录.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) // 1024}KB')
